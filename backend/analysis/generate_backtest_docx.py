#!/usr/bin/env python3
"""Generate DOCX report from brazil basis backtest results."""

from __future__ import annotations

import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
RESULTS = Path(__file__).resolve().parent / "data" / "backtest_results.json"
OUT_DOCX = ROOT / "docs" / "brazil-basis-backtest.docx"


def cents(v: float) -> str:
    return f"{v * 100:.2f}¢/lb"


def pct(v: float) -> str:
    return f"{v:.1f}%"


def run_backtest() -> None:
    subprocess.run(
        [sys.executable, str(Path(__file__).parent / "brazil_basis_backtest.py")],
        check=True,
        cwd=Path(__file__).resolve().parents[1],
    )


def load_results() -> dict:
    if not RESULTS.exists():
        run_backtest()
    return json.loads(RESULTS.read_text(encoding="utf-8"))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build_doc(data: dict) -> Document:
    doc = Document()
    section = doc.sections[0]

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("cotton.xyz\n")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x8B, 0x56, 0xE9)
    sub = title.add_run("Brazil Basis Model — Backtest Report\n")
    sub.bold = True
    sub.font.size = Pt(16)
    meta = title.add_run(
        f"BRAZIL/USDC pair · FX-linked ICE anchor model\n"
        f"Generated {datetime.utcnow().strftime('%Y-%m-%d')} · TEST / research\n"
    )
    meta.font.size = Pt(11)
    meta.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    add_heading(doc, "1. Executive Summary", 1)
    rating = data.get("rating_recommended") or data["rating"]
    version = data.get("version", 1)
    recommended = data.get("recommended_model", "hybrid_oracle_504d")
    doc.add_paragraph(
        f"This report evaluates the cotton.xyz BRAZIL/USDC mark engine (v{version}): "
        f"ICE relayer mark plus governance structural premium β_const."
    )
    p = doc.add_paragraph()
    p.add_run("Recommended model (holdout): ").bold = True
    p.add_run(f"{recommended}\n")
    p.add_run("Accuracy rating: ").bold = True
    grade_run = p.add_run(f"{rating['grade']} — {rating['score']}/100")
    grade_run.bold = True
    if rating["grade"] in ("A", "B"):
        grade_run.font.color.rgb = RGBColor(0x16, 0xA3, 0x34)
    elif rating["grade"] == "C":
        grade_run.font.color.rgb = RGBColor(0xCA, 0x8A, 0x04)
    else:
        grade_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    doc.add_paragraph(rating["summary"])

    holdout = data.get("models_holdout", {})
    if holdout:
        h = holdout.get("hybrid_oracle", {})
        ice = holdout.get("ice_only", {})
        v1 = holdout.get("static_v1", {})
        doc.add_paragraph(
            f"Holdout from {data.get('holdout_from', '2012')}: hybrid oracle MAE "
            f"{h.get('mae_basis', 0)*100:.2f}¢/lb vs static v1 "
            f"{v1.get('mae_basis', 0)*100:.2f}¢/lb vs ICE-only "
            f"{ice.get('mae_basis', 0)*100:.2f}¢/lb. "
            f"Hybrid uses lagged CEPEA when fresh (≤2 calendar-day gap) and "
            f"504-day trailing FX model on stale sessions (weekends/holidays)."
        )
    elif version == 1:
        doc.add_paragraph(
            "Key finding (v1): static FX + season model did not beat ICE-only on holdout."
        )

    add_heading(doc, "2. Model Specification", 1)
    doc.add_paragraph("Observed basis (ground truth):")
    doc.add_paragraph("Basis_obs = CEPEA_USD/lb − ICE_USD/lb", style="Intense Quote")
    doc.add_paragraph("Model (v3 — ICE + constant structural premium):")
    doc.add_paragraph(
        "Mark_Brazil = Mark_ICE + β_const\n"
        "β_const = median(CEPEA − ICE) over trailing 504 sessions\n"
        "Residual basis discovered on-book by traders",
        style="Intense Quote",
    )
    verdict = data.get("acceptance_verdict", {})
    if verdict:
        p = doc.add_paragraph()
        p.add_run("Acceptance vs actual CEPEA: ").bold = True
        p.add_run(verdict.get("summary", ""))
    doc.add_paragraph(
        "CEPEA is the ESALQ/USP cotton lint indicator (8-day payment); "
        "ICE is the NYBOT Cotton No. 2 future (CT=F proxy)."
    )

    add_heading(doc, "3. Data Sources", 1)
    src = data["data_sources"]
    add_table(
        doc,
        ["Series", "Source", "Notes"],
        [
            ["CEPEA algodão (8 dias)", "CEPEA via royopa/cepea_scraper CSV", src["cepea"]],
            ["ICE Cotton No. 2", src["ice"], "Front-month future; cents/lb → USD/lb"],
        ],
    )
    panel = data["panel"]
    doc.add_paragraph(
        f"Aligned sample: {panel['n']:,} trading days from {panel['start']} to {panel['end']}."
    )

    add_heading(doc, "4. Methodology", 1)
    meth = data.get("methodology", {})
    doc.add_paragraph(
        "• Holdout: final 30% of sample (2012–2018).\n"
        "• β_const: trailing 504-session median of (CEPEA_USD − ICE_USD).\n"
        "• Mark: ICE + β_const; residual basis left to traders on-book.\n"
        "• Baseline: ICE-only (basis = 0).\n"
        "• Metrics: MAE, RMSE, R²; % within 1%, ±2¢, ±5¢ of actual CEPEA mark."
    )

    add_heading(doc, "5. Production β_const (504d median)", 1)
    params = data.get("params", {})
    beta = params.get("beta_const_usd_lb") or data.get("production_beta_const_usd_lb", 0)
    doc.add_paragraph(f"β_const (structural premium): {beta * 100:.2f}¢/lb")
    doc.add_paragraph(f"Method: {params.get('method', 'median')} over {params.get('window_days', 504)} sessions")

    add_heading(doc, "6. Backtest Results", 1)
    holdout = data.get("models_holdout", {})
    if holdout:
        add_heading(doc, "Holdout comparison (recommended v2)", 2)
        rows = []
        for name, key in [
            ("ICE + rolling 504d median β (recommended)", "rolling_504d_median"),
            ("ICE + static β train median", "static_median_train"),
            ("ICE + static β train mean", "static_mean_train"),
            ("ICE-only", "ice_only"),
        ]:
            m = holdout.get(key, {})
            if not m:
                continue
            imp = ""
            if key != "ice_only" and holdout.get("ice_only"):
                ice_mae = holdout["ice_only"]["mae_basis"]
                imp = f"{((ice_mae - m['mae_basis']) / ice_mae * 100):+.1f}%"
            rows.append([
                name,
                cents(m["mae_basis"]),
                cents(m["rmse_basis"]),
                f"{m['r2_basis']:.2f}",
                pct(m["pct_within_5c"]),
                imp or "baseline",
            ])
        add_table(
            doc,
            ["Method", "MAE", "RMSE", "R²", "Within ±5¢", "vs ICE-only"],
            rows,
        )
    tm = data.get("test_metrics") or holdout.get("hybrid_oracle", {})
    ice = data.get("baselines", {}).get("ice_only") or holdout.get("ice_only", {})
    mean = data.get("baselines", {}).get("mean_basis", {})
    if tm and not holdout:
        add_table(
            doc,
            ["Method", "Test MAE", "RMSE", "R²", "Within ±2¢", "Within ±5¢", "vs ICE-only"],
            [
                [
                    "FX + season model",
                    cents(tm["mae_basis"]),
                    cents(tm["rmse_basis"]),
                    f"{tm['r2_basis']:.2f}",
                    pct(tm["pct_within_2c"]),
                    pct(tm["pct_within_5c"]),
                    f"{rating['vs_ice_only_improvement_pct']:+.1f}%",
                ],
                [
                    "ICE-only (basis = 0)",
                    cents(ice["mae_basis"]),
                    cents(ice["rmse_basis"]),
                    f"{ice['r2_basis']:.2f}",
                    pct(ice["pct_within_2c"]),
                    pct(ice["pct_within_5c"]),
                    "baseline",
                ],
            ],
        )

    wf = data.get("walk_forward")
    if wf:
        wfr = data.get("walk_forward_rating", {})
        add_heading(doc, "Walk-forward (by year)", 2)
        doc.add_paragraph(
            f"MAE {cents(wf['mae_basis'])}, RMSE {cents(wf['rmse_basis'])}, "
            f"R² {wf['r2_basis']:.2f}, grade {wfr.get('grade', '—')} "
            f"({wfr.get('score', '—')}/100)."
        )

    roll = data.get("rolling_expanding")
    if roll:
        rr = data.get("rolling_rating", {})
        add_heading(doc, "Rolling expanding refit", 2)
        doc.add_paragraph(
            f"One-step-ahead MAE {cents(roll['mae_basis'])}, "
            f"within ±5¢: {pct(roll['pct_within_5c'])}, "
            f"grade {rr.get('grade', '—')} ({rr.get('score', '—')}/100)."
        )

    corr = data.get("fx_basis_correlation")
    if corr:
        add_heading(doc, "FX vs basis correlation", 2)
        doc.add_paragraph(
            f"z(FX) vs basis_obs — full sample: {corr['full']:.3f}; "
            f"pre-2013: {corr['train_pre_2013']:.3f}; "
            f"2013+: {corr['test_2013_on']:.3f}. "
            "Negative sign matches theory (weaker BRL → higher USD basis) but "
            "correlation magnitude is modest."
        )

    add_heading(doc, "7. Accuracy Rating Scale", 1)
    add_table(
        doc,
        ["Grade", "Criteria (holdout test)"],
        [
            ["A", "MAE ≤ 2.5¢, R² ≥ 0.55, beats ICE-only materially"],
            ["B", "MAE ≤ 4¢, R² ≥ 0.35, beats ICE-only"],
            ["C", "Moderate improvement over ICE-only"],
            ["D", "Weak explanatory power"],
            ["F", "Does not beat ICE-only anchor"],
        ],
    )
    doc.add_paragraph(
        f"Holdout result: Grade {rating['grade']} ({rating['score']}/100)."
    )

    add_heading(doc, "8. Interpretation & Recommendations", 1)
    doc.add_paragraph(
        "Regime shift: train-period mean basis (~3.2¢/lb) vs test-period (~6.6¢/lb) "
        "indicates static parameters from 2000–2012 do not generalise to 2013–2018."
    )
    doc.add_paragraph(
        "Recommended path for cotton.xyz BRAZIL/USDC (testnet → live testnet):"
    )
    recs = [
        "Hybrid oracle: blend observed CEPEA when fresh (<24h); FX model as fallback only.",
        "Rolling recalibration: refit β coefficients monthly/quarterly on trailing 2–3 years.",
        "Extend CEPEA history beyond May 2018 for modern out-of-sample validation.",
        "Do not rely on FX + season model alone for production marks without CEPEA override.",
        "Keep US/USDC on pure ICE relayer mark; apply basis layer only on BRAZIL/USDC.",
    ]
    for r in recs:
        doc.add_paragraph(r, style="List Bullet")

    add_heading(doc, "9. Limitations", 1)
    for lim in data.get("limitations", []):
        doc.add_paragraph(lim, style="List Bullet")
    doc.add_paragraph(
        "CEPEA CSV ends 2018-05-08; recent spot checks (May 2026) show model bias "
        "when using stale calibration — observed basis ~+7¢ vs model ~−2.5¢ under "
        "current ICE/FX levels.",
        style="List Bullet",
    )

    add_heading(doc, "10. Reproduction", 1)
    doc.add_paragraph(
        "From the project repository:\n\n"
        "cd backend && source .venv/bin/activate\n"
        "python analysis/brazil_basis_backtest.py\n"
        "python analysis/generate_backtest_docx.py\n\n"
        f"JSON results: backend/analysis/data/backtest_results.json\n"
        f"Report output: docs/brazil-basis-backtest.docx"
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = footer.add_run("cotton.xyz · Confidential research · TEST environment")
    f.font.size = Pt(9)
    f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


def main() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    if not RESULTS.exists() or json.loads(RESULTS.read_text()).get("version", 1) < 3:
        print("Running backtest v3...")
        run_backtest()
    data = json.loads(RESULTS.read_text(encoding="utf-8"))
    doc = build_doc(data)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
